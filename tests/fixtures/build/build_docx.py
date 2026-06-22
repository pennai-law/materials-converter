"""Build tests/fixtures/sample_with_comments.docx — a Word document that exercises:
  - H1 + H2 headings (for section markers)
  - One footnote (preserved by default per spec §4.6)
  - One Word comment (dropped by default; appears with --full)
  - One tracked-changes pair (insertion + deletion; accepted-final by default)

python-docx builds the prose and headings. Comments are injected via lxml
since python-docx's high-level API doesn't support them. Tracked changes
likewise go in via lxml.

Run from repo root:
    ./venv/bin/python tests/fixtures/build/build_docx.py
"""
from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

# Safe XML parser: don't resolve external entities, don't make network requests.
# Mitigates XXE attacks where a crafted .docx contains DTD entity declarations
# pointing at file:// URLs or http:// servers.
_SAFE_XML_PARSER = etree.XMLParser(resolve_entities=False, no_network=True)

OUT = Path(__file__).resolve().parent.parent / "sample_with_comments.docx"

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _build_base_doc() -> Document:
    doc = Document()
    doc.add_heading("Memo Title", level=1)

    p1 = doc.add_paragraph(
        "This is the opening paragraph. It contains a phrase that a "
    )
    p1.add_run("contested claim")
    p1.add_run(" worth annotating, plus a footnote reference.")

    doc.add_heading("Discussion", level=2)
    doc.add_paragraph(
        "This paragraph contains both a footnote reference and a tracked "
        "change. The footnote text appears at the end of the document."
    )
    return doc


def _save_via_python_docx(doc: Document) -> None:
    doc.save(str(OUT))


def _inject_comment_into_archive() -> None:
    parts: dict[str, bytes] = {}
    with zipfile.ZipFile(OUT, "r") as zin:
        for name in zin.namelist():
            parts[name] = zin.read(name)

    comments_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:comment w:id="0" w:author="Sarah Chen" w:date="2025-09-12T10:00:00Z" w:initials="SC">'
        b'<w:p><w:r><w:t>Consider citing Bilski here.</w:t></w:r></w:p>'
        b'</w:comment>'
        b'</w:comments>'
    )
    parts["word/comments.xml"] = comments_xml

    doc_tree = etree.fromstring(parts["word/document.xml"], _SAFE_XML_PARSER)
    body = doc_tree.find(qn("w:body"))
    for paragraph in body.findall(qn("w:p")):
        for run in paragraph.findall(qn("w:r")):
            t = run.find(qn("w:t"))
            if t is not None and t.text == "contested claim":
                idx = list(paragraph).index(run)
                start = etree.SubElement(paragraph, qn("w:commentRangeStart"))
                start.set(qn("w:id"), "0")
                paragraph.remove(start)
                paragraph.insert(idx, start)
                end = etree.SubElement(paragraph, qn("w:commentRangeEnd"))
                end.set(qn("w:id"), "0")
                paragraph.remove(end)
                paragraph.insert(idx + 2, end)
                ref_run = etree.SubElement(paragraph, qn("w:r"))
                ref = etree.SubElement(ref_run, qn("w:commentReference"))
                ref.set(qn("w:id"), "0")
                paragraph.remove(ref_run)
                paragraph.insert(idx + 3, ref_run)
                break
        else:
            continue
        break
    parts["word/document.xml"] = etree.tostring(
        doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    rels_path = "word/_rels/document.xml.rels"
    rels_tree = etree.fromstring(parts[rels_path], _SAFE_XML_PARSER)
    R = "http://schemas.openxmlformats.org/package/2006/relationships"
    new_rel = etree.SubElement(rels_tree, f"{{{R}}}Relationship")
    new_rel.set("Id", "rIdComments")
    new_rel.set(
        "Type",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments",
    )
    new_rel.set("Target", "comments.xml")
    parts[rels_path] = etree.tostring(
        rels_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    ct_path = "[Content_Types].xml"
    ct_tree = etree.fromstring(parts[ct_path], _SAFE_XML_PARSER)
    CT = "http://schemas.openxmlformats.org/package/2006/content-types"
    override = etree.SubElement(ct_tree, f"{{{CT}}}Override")
    override.set("PartName", "/word/comments.xml")
    override.set(
        "ContentType",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml",
    )
    parts[ct_path] = etree.tostring(
        ct_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def _inject_tracked_change() -> None:
    parts: dict[str, bytes] = {}
    with zipfile.ZipFile(OUT, "r") as zin:
        for name in zin.namelist():
            parts[name] = zin.read(name)

    doc_tree = etree.fromstring(parts["word/document.xml"], _SAFE_XML_PARSER)
    body = doc_tree.find(qn("w:body"))
    paragraphs = body.findall(qn("w:p"))
    target = None
    for p in paragraphs:
        for r in p.findall(qn("w:r")):
            t = r.find(qn("w:t"))
            if t is not None and t.text and "tracked change" in t.text:
                target = p
                break
        if target is not None:
            break

    if target is None:
        return

    ins = etree.SubElement(target, qn("w:ins"))
    ins.set(qn("w:id"), "100")
    ins.set(qn("w:author"), "Sarah Chen")
    ins.set(qn("w:date"), "2025-09-12T10:00:00Z")
    ins_run = etree.SubElement(ins, qn("w:r"))
    ins_t = etree.SubElement(ins_run, qn("w:t"))
    ins_t.text = " ADDED-TEXT"
    ins_t.set(qn("xml:space"), "preserve")

    delete = etree.SubElement(target, qn("w:del"))
    delete.set(qn("w:id"), "101")
    delete.set(qn("w:author"), "Sarah Chen")
    delete.set(qn("w:date"), "2025-09-12T10:00:00Z")
    del_run = etree.SubElement(delete, qn("w:r"))
    del_t = etree.SubElement(del_run, qn("w:delText"))
    del_t.text = " REMOVED-TEXT"
    del_t.set(qn("xml:space"), "preserve")

    parts["word/document.xml"] = etree.tostring(
        doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def _inject_footnote() -> None:
    parts: dict[str, bytes] = {}
    with zipfile.ZipFile(OUT, "r") as zin:
        for name in zin.namelist():
            parts[name] = zin.read(name)

    footnotes_xml = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
        b'<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
        b'<w:footnote w:id="1">'
        b'<w:p><w:r><w:t>This is the footnote text. It contains a citation: Bilski v. Kappos, 561 U.S. 593 (2010).</w:t></w:r></w:p>'
        b'</w:footnote>'
        b'</w:footnotes>'
    )
    parts["word/footnotes.xml"] = footnotes_xml

    doc_tree = etree.fromstring(parts["word/document.xml"], _SAFE_XML_PARSER)
    body = doc_tree.find(qn("w:body"))
    paragraphs = body.findall(qn("w:p"))
    target = None
    for p in paragraphs:
        for r in p.findall(qn("w:r")):
            t = r.find(qn("w:t"))
            if t is not None and t.text and "footnote text appears" in t.text:
                target = p
                break
        if target is not None:
            break
    if target is None and paragraphs:
        target = paragraphs[-1]

    if target is not None:
        ref_run = etree.SubElement(target, qn("w:r"))
        ref = etree.SubElement(ref_run, qn("w:footnoteReference"))
        ref.set(qn("w:id"), "1")

    parts["word/document.xml"] = etree.tostring(
        doc_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    rels_path = "word/_rels/document.xml.rels"
    rels_tree = etree.fromstring(parts[rels_path], _SAFE_XML_PARSER)
    R = "http://schemas.openxmlformats.org/package/2006/relationships"
    new_rel = etree.SubElement(rels_tree, f"{{{R}}}Relationship")
    new_rel.set("Id", "rIdFootnotes")
    new_rel.set(
        "Type",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes",
    )
    new_rel.set("Target", "footnotes.xml")
    parts[rels_path] = etree.tostring(
        rels_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    ct_path = "[Content_Types].xml"
    ct_tree = etree.fromstring(parts[ct_path], _SAFE_XML_PARSER)
    CT = "http://schemas.openxmlformats.org/package/2006/content-types"
    override = etree.SubElement(ct_tree, f"{{{CT}}}Override")
    override.set("PartName", "/word/footnotes.xml")
    override.set(
        "ContentType",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml",
    )
    parts[ct_path] = etree.tostring(
        ct_tree, xml_declaration=True, encoding="UTF-8", standalone=True,
    )

    with zipfile.ZipFile(OUT, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in parts.items():
            zout.writestr(name, data)


def build() -> Path:
    doc = _build_base_doc()
    _save_via_python_docx(doc)
    _inject_comment_into_archive()
    _inject_tracked_change()
    _inject_footnote()
    return OUT


if __name__ == "__main__":
    p = build()
    size_kb = p.stat().st_size / 1024
    print(f"Wrote {p} ({size_kb:.1f} KB)")
